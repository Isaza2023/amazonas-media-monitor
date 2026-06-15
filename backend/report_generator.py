import io
import json
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def generate_csv_report(articles) -> bytes:
    """Genera un archivo CSV con las noticias seleccionadas."""
    df = pd.DataFrame(articles)
    output = io.StringIO()
    df.to_csv(output, index=False, encoding='utf-8-sig')
    return output.getvalue().encode('utf-8-sig')

def generate_json_report(articles) -> bytes:
    """Genera un archivo JSON con las noticias seleccionadas."""
    return json.dumps(articles, indent=4, default=str, ensure_ascii=False).encode('utf-8')

def generate_excel_report(articles) -> bytes:
    """Genera un archivo Excel con las noticias seleccionadas."""
    df = pd.DataFrame(articles)
    # Formatear fechas para Excel si existen
    if 'publish_date' in df.columns:
        df['publish_date'] = df['publish_date'].apply(
            lambda x: x.strftime('%Y-%m-%d %H:%M:%S') if isinstance(x, datetime) else str(x)
        )
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Monitoreo_Amazonas')
    return output.getvalue()

def generate_docx_report(articles) -> bytes:
    """Genera el Reporte Flash en formato Word (DOCX) según la plantilla institucional."""
    doc = Document()
    
    # Configurar márgenes estándar
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72) # 1 pulgada
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # Estilos del encabezado
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_header1 = p_header.add_run("DIRECCIÓN SECCIONAL AMAZONAS\n")
    r_header1.font.bold = True
    r_header1.font.size = Pt(12)
    r_header1.font.name = 'Arial'
    r_header1.font.color.rgb = RGBColor(16, 44, 87) # Azul institucional

    r_header2 = p_header.add_run("SECCIÓN DE POLICÍA JUDICIAL CTI AMAZONAS\n")
    r_header2.font.bold = True
    r_header2.font.size = Pt(11)
    r_header2.font.name = 'Arial'
    r_header2.font.color.rgb = RGBColor(100, 110, 120)

    r_title = p_header.add_run("\nREPORTE FLASH DE MONITOREO DE MEDIOS\n")
    r_title.font.bold = True
    r_title.font.size = Pt(14)
    r_title.font.name = 'Arial'
    r_title.font.color.rgb = RGBColor(16, 44, 87)
    
    # Línea horizontal divisoria
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("__________________________________________________________________\n")
    r_sep.font.color.rgb = RGBColor(200, 200, 200)

    # Si no hay noticias
    if not articles:
        p = doc.add_paragraph()
        p.add_run("No se seleccionaron noticias para este reporte.").font.italic = True
    else:
        for index, art in enumerate(articles, 1):
            # Título de sección para cada noticia
            p_num = doc.add_paragraph()
            r_num = p_num.add_run(f"CASO / NOTICIA N° {index}: {art.get('title', 'Sin Título')}\n")
            r_num.font.bold = True
            r_num.font.size = Pt(11)
            r_num.font.color.rgb = RGBColor(16, 44, 87)

            # Diccionario con los datos a rellenar
            fields = [
                ("Fecha y hora del reporte", datetime.now().strftime('%d/%m/%Y %H:%M:%S')),
                ("Periodo monitoreado", "Últimas 24 Horas" if index == 1 else "Continuo"),
                ("Tema principal", art.get("topic", "General")),
                ("Nivel de relevancia", art.get("relevance", "BAJA")),
                ("Fuente", art.get("source", "No especificada")),
                ("Link original", art.get("url", "No especificado")),
                ("Lugar relacionado", art.get("location", "Amazonas")),
                ("Resumen", art.get("summary", "Sin resumen")),
                ("Descripción de la publicación o noticia", art.get("summary", "Ver link original")),
                ("Análisis preliminar", art.get("analysis_preliminary", "En proceso de análisis por el analista.")),
                ("Posible afectación para el Amazonas colombiano", art.get("regional_impact", "Bajo monitoreo preventivo por parte de las autoridades locales.")),
                ("Palabras clave detectadas", art.get("keywords", "Amazonas, Leticia")),
                ("Observaciones del analista", art.get("analyst_notes", "Sin observaciones adicionales.")),
                ("Elaboró", art.get("elaborated_by", "Analista de Medios CTI"))
            ]

            for label, value in fields:
                p_field = doc.add_paragraph()
                p_field.paragraph_format.left_indent = Pt(18)
                p_field.paragraph_format.space_after = Pt(2)
                p_field.paragraph_format.space_before = Pt(2)
                
                # Etiqueta
                r_label = p_field.add_run(f"{label}: ")
                r_label.font.bold = True
                r_label.font.size = Pt(10)
                r_label.font.name = 'Arial'
                r_label.font.color.rgb = RGBColor(40, 40, 40)
                
                # Valor
                r_val = p_field.add_run(f"{value}")
                r_val.font.size = Pt(10)
                r_val.font.name = 'Arial'
                if label == "Nivel de relevancia":
                    r_val.font.bold = True
                    if value == "ALTA":
                        r_val.font.color.rgb = RGBColor(180, 40, 40)  # Rojo
                    elif value == "MEDIA":
                        r_val.font.color.rgb = RGBColor(220, 160, 40) # Naranja
                    else:
                        r_val.font.color.rgb = RGBColor(40, 140, 40)  # Verde
                elif label == "Link original":
                    r_val.font.underline = True
                    r_val.font.color.rgb = RGBColor(0, 0, 238) # Azul enlace

            # Línea divisoria entre noticias
            if index < len(articles):
                p_div = doc.add_paragraph()
                r_div = p_div.add_run("\n" + "-"*40 + "\n")
                r_div.font.color.rgb = RGBColor(220, 220, 220)

    # Guardar en buffer de bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()


class FlashReportPDF(FPDF):
    """Generador de PDF profesional con FPDF2."""
    def header(self):
        # Dibujar una franja superior azul institucional
        self.set_fill_color(16, 44, 87)
        self.rect(0, 0, 210, 12, 'F')
        
        self.set_y(18)
        self.set_font('Helvetica', 'B', 11)
        self.set_text_color(16, 44, 87)
        self.cell(0, 5, 'DIRECCIÓN SECCIONAL AMAZONAS', 0, 1, 'C')
        
        self.set_font('Helvetica', 'B', 9)
        self.set_text_color(100, 110, 120)
        self.cell(0, 5, 'SECCIÓN DE POLICÍA JUDICIAL CTI AMAZONAS', 0, 1, 'C')
        
        self.set_font('Helvetica', 'B', 12)
        self.set_text_color(16, 44, 87)
        self.cell(0, 7, 'REPORTE FLASH DE MONITOREO DE MEDIOS', 0, 1, 'C')
        
        # Línea de separación
        self.set_draw_color(200, 200, 200)
        self.line(15, 38, 195, 38)
        self.ln(6)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f'Página {self.page_no()}', 0, 0, 'C')


def sanitize_pdf_text(text: str) -> str:
    """Reemplaza caracteres especiales fuera de Latin-1 para evitar excepciones en FPDF2."""
    if not text:
        return ""
    # Mapeo de reemplazos para caracteres tipográficos comunes que causan error
    replacements = {
        '\u2013': '-', # en-dash
        '\u2014': '-', # em-dash
        '\u2018': "'", # comilla simple izquierda
        '\u2019': "'", # comilla simple derecha
        '\u201c': '"', # comilla doble izquierda
        '\u201d': '"', # comilla doble derecha
        '\u2026': '...', # elipsis (...)
        '\u00a0': ' ',   # espacio duro
        '\u200b': '',    # espacio de ancho cero
        '\u20ac': 'EUR', # euro
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Codificar a latin-1 ignorando o reemplazando caracteres no mapeados
    return text.encode('latin-1', 'replace').decode('latin-1')


def generate_pdf_report(articles) -> bytes:
    """Genera el Reporte Flash en formato PDF."""
    pdf = FlashReportPDF()
    pdf.set_margins(15, 20, 15)
    pdf.alias_nb_pages()
    pdf.add_page()
    
    if not articles:
        pdf.set_font('Helvetica', 'I', 10)
        pdf.cell(0, 10, 'No se seleccionaron noticias para este reporte.', 0, 1, 'L')
    else:
        for index, art in enumerate(articles, 1):
            pdf.set_fill_color(245, 247, 250)
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_text_color(16, 44, 87)
            
            # Encabezado del caso
            title_text = sanitize_pdf_text(f"CASO / NOTICIA N° {index}: {art.get('title', 'Sin Título')}")
            pdf.multi_cell(0, 6, title_text, 0, 'L', fill=True)
            pdf.ln(2)

            # Campos del Reporte Flash
            fields = [
                ("Fecha y hora del reporte", datetime.now().strftime('%d/%m/%Y %H:%M:%S')),
                ("Periodo monitoreado", "Últimas 24 Horas" if index == 1 else "Continuo"),
                ("Tema principal", art.get("topic", "General")),
                ("Nivel de relevancia", art.get("relevance", "BAJA")),
                ("Fuente", art.get("source", "No especificada")),
                ("Link original", art.get("url", "No especificado")),
                ("Lugar relacionado", art.get("location", "Amazonas")),
                ("Resumen", art.get("summary", "Sin resumen")),
                ("Descripción de la publicación o noticia", art.get("summary", "Ver link original")),
                ("Análisis preliminar", art.get("analysis_preliminary", "En proceso de análisis por el analista.")),
                ("Posible afectación para el Amazonas colombiano", art.get("regional_impact", "Bajo monitoreo preventivo por parte de las autoridades locales.")),
                ("Palabras clave detectadas", art.get("keywords", "Amazonas, Leticia")),
                ("Observaciones del analista", art.get("analyst_notes", "Sin observaciones adicionales.")),
                ("Elaboró", art.get("elaborated_by", "Analista de Medios CTI"))
            ]

            for label, value in fields:
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(50, 50, 50)
                pdf.write(5, f"  {label}: ")
                
                # Formatear el valor según el campo
                pdf.set_font('Helvetica', '', 9)
                if label == "Nivel de relevancia":
                    pdf.set_font('Helvetica', 'B', 9)
                    if value == "ALTA":
                        pdf.set_text_color(180, 40, 40)
                    elif value == "MEDIA":
                        pdf.set_text_color(220, 160, 40)
                    else:
                        pdf.set_text_color(40, 140, 40)
                elif label == "Link original":
                    pdf.set_text_color(0, 0, 238) # Azul
                else:
                    pdf.set_text_color(80, 80, 80)
                
                # Procesar cadenas largas con multi_cell si exceden línea
                # write dibuja inline, si es un texto largo usamos multi_cell
                val_str = sanitize_pdf_text(str(value))
                if len(val_str) > 70 or "\n" in val_str:
                    pdf.ln(5)
                    pdf.set_x(20) # sangría
                    pdf.multi_cell(0, 4.5, val_str, 0, 'L')
                else:
                    pdf.write(5, f"{val_str}\n")
                
                # Restaurar espaciado
                pdf.set_text_color(50, 50, 50)
            
            pdf.ln(4)
            # Dibujar línea divisoria si quedan más casos
            if index < len(articles):
                pdf.set_draw_color(230, 230, 230)
                pdf.line(15, pdf.get_y(), 195, pdf.get_y())
                pdf.ln(4)
                
                # Si estamos cerca del fondo de la página, añadir página
                if pdf.get_y() > 240:
                    pdf.add_page()
                    
    return bytes(pdf.output())
